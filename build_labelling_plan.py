from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

OUT = Path('AIOps_Hand_Gaze_Step_Label_Assist_Implementation_Plan.docx')
NAVY = RGBColor(11, 37, 69); BLUE = RGBColor(46, 116, 181); GRAY = RGBColor(85,85,85)
LIGHT = 'E8EEF5'; PALE = 'F4F6F9'; RED='FCE8E6'; GREEN='E6F4EA'; GOLD='FFF4CE'

doc = Document()
sec = doc.sections[0]
sec.top_margin=Inches(0.75); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85)
sec.header_distance=Inches(0.35); sec.footer_distance=Inches(0.35)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor(25,25,25)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for nm,size,col,bef,aft in [('Title',28,NAVY,0,8),('Heading 1',17,BLUE,14,7),('Heading 2',13.5,BLUE,11,5),('Heading 3',11.5,NAVY,8,3)]:
    s=styles[nm]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=col; s.font.bold=(nm!='Title'); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True

header=sec.header.paragraphs[0]; header.text='AIOps Dataset Labelling | Technical Implementation Plan'; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.size=Pt(8.5); r.font.color.rgb=GRAY
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=footer.add_run('Implementation-ready research plan | 16 July 2026'); run.font.size=Pt(8); run.font.color.rgb=GRAY

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)
def margins(cell, top=90, start=110, bottom=90, end=110):
    tc=cell._tc.get_or_add_tcPr(); mar=tc.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); tc.append(mar)
    for side,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=mar.find(qn('w:'+side))
        if el is None: el=OxmlElement('w:'+side); mar.append(el)
        el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')
def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,LIGHT); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(c)
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.7)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t
def callout(label, text, fill=PALE):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.65)
    c=t.cell(0,0); shade(c,fill); margins(c,120,150,120,150)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(label+'  '); r.bold=True; r.font.color.rgb=NAVY; p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
def bullets(items, style='List Bullet'):
    for x in items: doc.add_paragraph(x, style=style)
def numbered(items):
    for x in items: doc.add_paragraph(x, style='List Number')
def linktext(text,url):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.add_run(text+' — ').bold=True; p.add_run(url)

# cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(32); p.add_run('TECHNICAL FIELD GUIDE').bold=True; p.runs[0].font.color.rgb=BLUE; p.runs[0].font.size=Pt(11)
doc.add_paragraph('Hand Joints, Gaze and Procedure-Step Label Assist', style='Title')
p=doc.add_paragraph('A fool-proof, human-in-the-loop plan for a small deployable model, with optional tool identity and 6D orientation.'); p.runs[0].font.size=Pt(14); p.runs[0].font.color.rgb=GRAY
doc.add_paragraph('Prepared for the AIOps dataset collection project | Version 1.0 | 16 July 2026')
callout('Recommendation', 'Build a modular offline pre-annotation service around CVAT. Use a pretrained 21-keypoint hand model, sensor-derived gaze when possible, and a lightweight temporal segmenter trained on frozen visual + hand features. Do not train one monolithic model and do not accept predictions without human review.', GREEN)
doc.add_paragraph('Decision summary', style='Heading 1')
table(['Capability','MVP decision','Why'],[
('Hand joints','MediaPipe Hand Landmarker proposals; fine-tune only after an error audit','Fast, 21 landmarks/hand, tracking-friendly; minimizes initial training.'),
('Gaze','Record calibrated eye-gaze ray; project to RGB/depth. RGB-only output is explicitly a proxy.','True wearer gaze is not observable from a forward-facing camera.'),
('Step','Frozen frame encoder + hand/object features → MS-TCN++ baseline; grammar/Viterbi post-processing.','Small-data friendly and temporally stable.'),
('Tool labels','Detector/segmenter trained after taxonomy freeze; track across frames.','High value but sensitive to class definitions and occlusion.'),
('Orientation','Defer to RGB-D + CAD/reference-image FoundationPose pilot.','6D pose requires camera calibration and object geometry/reference views.')], [1.05,2.25,3.35])

doc.add_page_break()
doc.add_paragraph('1. Scope, assumptions and non-negotiable definitions', style='Heading 1')
doc.add_paragraph('The target resembles IndustReal: egocentric assembly video with overlaid hand joints and gaze, plus a time-aligned procedure label. The system’s purpose is to reduce labelling time while preserving auditable ground truth.')
callout('Hard boundary', 'A model proposal is not ground truth. Store model version, confidence, and review status for every generated annotation. Ground truth exists only after review or a validated sensor measurement.', GOLD)
table(['Term','Operational definition'],[
('Hand joint','One of 21 anatomical landmarks per hand, in image pixels plus visibility; optional normalized image and camera/world coordinates.'),
('Gaze point','Intersection of a calibrated eye-gaze ray with the scene/depth surface, then projected into the RGB frame. Include validity and uncertainty radius.'),
('Gaze proxy','Predicted attention hotspot from scene context when no eye sensor exists. Never call this measured gaze.'),
('Step','A closed temporal interval [start, end) with exactly one canonical procedure-state label; background/transition is a valid class.'),
('Tool label','Tracked instance identity and class; “active” means involved in hand contact, not merely visible.'),
('Orientation','Object-to-camera rotation and translation, with coordinate convention, symmetry group, and pose validity.')],[1.35,5.3])

doc.add_paragraph('2. Acquisition design: decisions before collecting data', style='Heading 1')
bullets([
'Preferred rig: synchronized RGB (≥1080p/30 fps), depth when orientation is planned, eye-gaze ray at native rate, camera intrinsics/extrinsics, timestamps from one monotonic clock, and optional IMU.',
'Mount and field of view must be fixed within a collection phase. Recalibrate whenever camera, headset fit, resolution, or lens settings change.',
'Per participant/session: eye calibration, 9-point validation, hand visibility test, tool pose validation, clock drift check, and a 30-second neutral/background segment.',
'Capture variation deliberately: users, handedness, gloves/skin tones, lighting, benches, camera tilt, tool instances, partial occlusion, errors, step order deviations, and pauses.',
'Privacy: obtain explicit consent for eye/biometric signals; minimize identity metadata; encrypt raw streams; define retention and deletion before collection.'
])
callout('Go/no-go for gaze', 'If there is no eye tracker, revise the schema and UI to “attention proxy.” Do not spend the MVP budget training a face-based gaze model: the wearer’s face/eyes are absent in egocentric RGB.', RED)

doc.add_paragraph('3. Annotation contract and storage schema', style='Heading 1')
doc.add_paragraph('Use a versioned JSON/JSONL canonical store; export COCO keypoints for hand training and a frame/segment table for temporal models. Never make CVAT export files the sole source of truth.')
table(['Entity','Required fields'],[
('frame','video_id, frame_id, timestamp_ns, width, height, camera_id, calibration_id'),
('hand','track_id, side {L,R,unknown}, 21×(x_px,y_px,visibility), bbox, occluded, source, model_version, confidence, review_status'),
('gaze','origin_xyz, direction_xyz, point_xyz?, point_xy?, valid, uncertainty_deg/radius_px, calibration_id, source {sensor,proxy,manual}'),
('step segment','step_id, start_ns, end_ns, status {correct,incorrect,aborted,unknown}, prerequisites, annotator, review_status'),
('tool','track_id, class_id, bbox/mask, active, contact_hand, occlusion, source, confidence'),
('pose','track_id, R_3x3 or quaternion, t_xyz_m, coordinate_frame, symmetry_id, valid, reprojection_error_px')],[1.35,5.3])
doc.add_paragraph('Annotation rules that prevent ambiguity', style='Heading 2')
bullets([
'Joints: label visible anatomical positions; for occluded joints use visibility=0 and either omit coordinates or mark inferred coordinates separately. Do not silently hallucinate.',
'Handedness: identify from wearer perspective. Permit “unknown” rather than forcing a side.',
'Gaze: suppress invalid samples during blinks/tracking loss; resample only after timestamp alignment; never interpolate across long invalid gaps.',
'Step boundary: start at the first task-directed motion/contact specified in the guide; end at the first frame the defined completion condition becomes true. Keep completion-state and action labels separate.',
'Errors: encode incorrect completion and aborted attempts explicitly; do not coerce them into the nearest successful step.',
'Double review every new class and every ambiguity; maintain a living adjudication log with visual examples.'
])

doc.add_page_break()
doc.add_paragraph('4. System architecture', style='Heading 1')
doc.add_paragraph('Pipeline: ingest and synchronize → frame sampling → proposal models → temporal smoothing/constraints → CVAT review → canonical export → QA sampling → training set registry. Each module is replaceable and independently evaluated.')
table(['Service','Input','Output','MVP implementation'],[
('Synchronizer','RGB/depth/gaze/IMU timestamps','aligned sample index','Nearest timestamp with maximum-gap rule; log drift and dropped frames.'),
('Hand proposer','RGB frames','21 keypoints + side + confidence','MediaPipe Hand Landmarker; track IDs via temporal association; One-Euro/Kalman smoothing only for display.'),
('Gaze projector','gaze ray + calibration + depth/mesh','2D/3D point + uncertainty','Ray/depth intersection; calibrated transform; invalidate off-frame/no-hit samples.'),
('Feature extractor','RGB clips/crops','frozen embeddings','Small pretrained video/image encoder; cache features once.'),
('Step segmenter','embeddings + hand/tool signals','frame logits + boundaries','MS-TCN++ baseline; Viterbi/finite-state grammar constrained by procedure.'),
('Review adapter','proposals','CVAT shapes/tracks/tags','Local CVAT auto-annotation function and SDK import/export.'),
('Registry','models, labels, splits','reproducible dataset release','DVC/MLflow-style manifests or simple immutable object-store hashes.')],[1.15,1.3,1.55,2.65])

doc.add_paragraph('5. Model specifications', style='Heading 1')
doc.add_paragraph('5.1 Hand joints (MVP)', style='Heading 2')
numbered([
'Run a pretrained 21-landmark hand detector on every frame or every 2nd frame; use its tracking mode between detections.',
'Reject proposals below a tuned presence threshold; expand hand crops 20–30% to reduce fingertip truncation.',
'Associate tracks with handedness, bbox IoU, and landmark distance; allow track reset after prolonged occlusion.',
'Apply temporal smoothing to display/proposals, but retain raw coordinates. Do not smooth across cuts or missing tracks.',
'Fine-tune a small keypoint network only if the first 2,000 reviewed frames show systematic domain error (gloves, tools covering fingers, fisheye lens). Use corrected crops as training data.'
])
doc.add_paragraph('Suggested fine-tune: 256×256 hand crops, 21 heatmaps plus visibility, geometric/color augmentation, synthetic occluders shaped like tools, AdamW, early stopping on participant-held-out PCK. Start with 1–3k diverse corrected hands; expand only if the learning curve is still rising.')

doc.add_paragraph('5.2 Gaze (sensor-first)', style='Heading 2')
numbered([
'Calibrate each wearer; store gaze ray origin/direction with the device timestamp and calibration ID.',
'Transform gaze ray into the RGB camera coordinate frame using time-specific headset pose and fixed extrinsics.',
'Intersect with depth/scene mesh. If depth is unavailable, project the ray to an agreed working-plane estimate and mark depth_assumed=true.',
'Project through camera intrinsics/distortion to image pixels; attach uncertainty derived from angular accuracy, distance, calibration residual, and synchronization error.',
'Filter blinks/invalid tracking. Use a short robust median filter only for visualization; retain raw samples.',
'Validate each session against known targets and fail collection if median angular/point error exceeds the project threshold.'
])
callout('Practical threshold', 'For small targets, use a visible uncertainty circle rather than a cross alone. Microsoft describes HoloLens 2 gaze at ~30 Hz and approximately 1.5° around the target, recommending larger margins in interaction design; your measured session validation overrides this nominal figure.', GOLD)

doc.add_paragraph('5.3 Procedure-step segmentation (small model)', style='Heading 2')
doc.add_paragraph('Use a two-stage design. Stage A caches frozen visual embeddings at 5–10 Hz and concatenates compact cues: normalized hand joints/velocities, hand presence, gaze-to-tool distances, active-tool logits, and optional audio. Stage B is a small MS-TCN++ temporal model with 2–4 stages, 8–10 dilated residual layers, 64–128 channels, dropout, and framewise cross-entropy plus temporal smoothing loss.')
bullets([
'Train on whole sequences or long windows with overlap; class-balance without destroying real background frequency.',
'Post-process with minimum-duration rules and a procedure finite-state graph. Allow legal skips, repeats, errors, and rollback states—never enforce one ideal sequence.',
'Cold start: label 10–15 complete sessions across ≥5 people manually; freeze a participant-held-out validation/test split before tuning.',
'Active learning: rank segments by entropy, boundary disagreement, illegal transition, and model-versus-annotator disagreement; review these first.',
'Do not optimize frame accuracy alone; it rewards background and misses fragmented segments.'
])

doc.add_paragraph('5.4 Tool identity and orientation (aspirational)', style='Heading 2')
doc.add_paragraph('Tool labels: begin with 5–15 operationally meaningful classes after the taxonomy is frozen. Train a small detector/segmenter from boxes or masks and propagate with tracking. Store instance and class separately; add “other/unknown.” Orientation: use RGB-D, calibrated intrinsics, per-tool CAD or reference views, and FoundationPose. Define object frames and symmetries before annotation. Report pose invalid when the object is mostly occluded or symmetric orientation is unobservable.')

doc.add_page_break()
doc.add_paragraph('6. Training, splits and leakage control', style='Heading 1')
table(['Rule','Required implementation'],[
('Split unit','Participant × session × physical setup; never random frames. Keep all frames from a sequence in one split.'),
('Test freeze','Freeze test IDs and labels before the first model comparison. No threshold tuning on test.'),
('Versioning','Hash raw videos, calibration, label schema, taxonomy, split manifest, code commit, weights and inference config.'),
('Pretraining','Record datasets/licenses and avoid test-environment leakage where possible.'),
('Seed/repeat','At least 3 seeds for the temporal model if reporting comparative results.'),
('Failure slices','Gloves, skin tone, handedness, lighting, occlusion, fast motion, tool class, participant, error steps, short steps.')],[1.3,5.35])

doc.add_paragraph('7. Evaluation and acceptance gates', style='Heading 1')
table(['Capability','Metrics','Pilot acceptance gate (set after baseline)'],[
('Hand joints','PCK@0.05/0.1 normalized by hand bbox; mean pixel error; visibility F1; track ID switches','≥95% hands detected on in-domain review set; PCK@0.1 ≥0.90; no subgroup >10-point gap without remediation.'),
('Gaze','valid-sample rate; angular error; pixel/3D point error; target-hit accuracy; calibration drift','Session validation passes predefined target error; ≥90% valid during usable footage; uncertainty displayed.'),
('Steps','segmental F1@10/25/50; edit score; boundary MAE; per-class recall; frame accuracy secondary','F1@50 ≥0.75 on held-out participants and boundary MAE ≤1 s for MVP; thresholds adjustable to task cadence.'),
('Tool','mAP50-95, mask AP, active-tool F1, track IDF1','Per-class recall ≥0.85 for safety/critical tools before auto-proposal is enabled.'),
('Orientation','ADD(-S), 2D reprojection, translation/rotation error, tracking loss','Gate per tool and symmetry; no aggregate-only acceptance.'),
('Productivity','median review seconds/frame or minute; acceptance rate; correction distance; escaped-error audit','≥2× reviewed-label throughput versus manual baseline with equal or better audited error rate.')],[1.05,2.3,3.3])
callout('Quality sampling', 'Blind-audit at least 10% of accepted proposals during pilot, stratified by confidence and failure slice. A second reviewer must not see whether a label was model-generated. Track false acceptance, not just correction rate.', RED)

doc.add_paragraph('8. Implementation roadmap (12 weeks)', style='Heading 1')
table(['Weeks','Deliverable','Exit criterion'],[
('0–1','Taxonomy, annotation guide, privacy review, rig decision, calibration protocol','20 representative clips adjudicated; no unresolved definitions.'),
('2','Capture/clock/calibration validation; canonical schema and split manifest','A 30-min pilot replays with <1-frame RGB alignment and documented gaze residuals.'),
('3–4','CVAT project; hand proposal adapter; gaze projection overlay','Annotator can accept/edit joints and gaze; provenance exports correctly.'),
('5','Manual gold set: 10–15 full sessions + double-labelled subset','Inter-annotator agreement meets thresholds or guide is revised.'),
('6–7','Frozen feature extraction + MS-TCN++ baseline + grammar','Reproducible training; held-out metrics and failure slices reported.'),
('8','Active-learning queue and confidence calibration','Uncertainty ranking finds more errors than random sampling.'),
('9','End-to-end productivity trial versus manual baseline','≥2× throughput without higher blinded-audit error.'),
('10–11','Fix dominant failure modes; package local inference service','One-command/container deployment; deterministic export.'),
('12','MVP release, model card, dataset card, operating SOP','All acceptance gates passed; rollback and monitoring tested.'),
('Later','Tool detector; RGB-D/CAD orientation pilot','Separate go/no-go evaluation; does not block MVP.')],[0.7,3.0,2.95])

doc.add_paragraph('9. Engineering backlog and repository layout', style='Heading 1')
doc.add_paragraph('Recommended structure: configs/ (taxonomy, calibration, splits), data_manifests/, src/sync, src/hands, src/gaze, src/steps, src/tools, src/cvat_adapter, tests/, evaluation/, model_cards/, annotation_guides/. Keep raw media outside Git; commit immutable manifests and checksums.')
bullets([
'CLI: ingest, validate-calibration, propose-hands, project-gaze, extract-features, train-steps, infer, export-cvat, import-reviews, evaluate, audit.',
'Unit tests: coordinate transforms, distortion round-trip, timestamp maximum-gap behavior, schema validation, handedness mapping, interval boundaries.',
'Golden integration test: a 30-second clip with fixed expected hashes and metric outputs.',
'Operational safeguards: model-version pinning, low-confidence suppression, per-module kill switch, resumable jobs, structured logs, and no overwrite of reviewed labels.',
'Hardware target: start with one 8–12 GB NVIDIA GPU for offline proposals; CPU-only hand inference is feasible but benchmark on actual footage. Storage dominates compute—budget for original streams plus derivatives and backups.'
])

doc.add_paragraph('10. Risks and mitigations', style='Heading 1')
table(['Risk','Detection','Mitigation'],[
('Egocentric hand occlusion','Low visibility, fingertip error near tools','Explicit occlusion labels; tool-shaped augmentation; crop margins; human review.'),
('Gaze presented as truth without sensor','No eye stream/calibration ID','Rename to proxy; prohibit merge into measured-gaze field.'),
('Step label leakage','Same participant/session in train and test','Manifest validator rejects overlapping groups.'),
('Over-segmentation','High frame accuracy, low edit/F1','Temporal smoothing loss, segment metrics, minimum-duration/grammar.'),
('Taxonomy churn','Frequent class merges/splits','Pilot taxonomy first; stable IDs; migration scripts.'),
('Automation bias','High acceptance but poor blind audit','Hide confidence for audit subset; second reviewer; confidence-stratified audit.'),
('Orientation symmetry','Large rotation error but identical appearance','Declare symmetry group; use ADD-S/symmetry-aware scoring; invalid state.'),
('Domain drift','Performance falls by user/tool/lighting','Slice dashboards; periodic gold set; retrain only after drift trigger.')],[1.35,2.0,3.3])

doc.add_page_break()
doc.add_paragraph('11. Pilot runbook', style='Heading 1')
numbered([
'Freeze schema v0.1, step graph, tool list, coordinate conventions, and 20 visual examples per ambiguous label.',
'Record three dry-run sessions; validate timestamps, calibration transforms, gaze targets, hand visibility, storage, and consent flow before recruiting broadly.',
'Create a manual gold set with two annotators; adjudicate disagreements and update the guide once.',
'Benchmark manual-only annotation time and blind-audit error.',
'Run pretrained hand proposals and sensor gaze projection; measure proposal acceptance and correction distance.',
'Train the temporal model only after enough complete sequences exist; freeze participant-held-out splits.',
'Run a randomized crossover productivity study: manual-first vs assisted-first clips, balanced across annotators.',
'Pass acceptance gates, document known failure slices, and release the MVP. Otherwise disable the failing module while keeping the rest.'
])
doc.add_paragraph('Definition of done', style='Heading 2')
bullets([
'One command turns synchronized raw streams into a reviewable CVAT task.',
'Every proposal has provenance, confidence, and immutable model/config identifiers.',
'Reviewed annotations round-trip without coordinate or timing loss.',
'Gold-set and productivity reports are reproducible from a frozen manifest.',
'Measured gaze and inferred attention proxy cannot be confused in schema or UI.',
'The system fails safely: missing calibration, invalid gaze, low-confidence hands, and illegal step transitions are surfaced—not silently filled.'
])

doc.add_paragraph('12. Research basis and citations', style='Heading 1')
doc.add_paragraph('Primary and authoritative sources used to shape the plan:')
sources=[
('Schoonbeek et al., “IndustReal: A Dataset for Procedure Step Recognition Handling Execution Errors…” WACV 2024.','https://openaccess.thecvf.com/content/WACV2024/papers/Schoonbeek_IndustReal_A_Dataset_for_Procedure_Step_Recognition_Handling_Execution_Errors_WACV_2024_paper.pdf'),
('IndustReal project and released annotation formats.','https://timschoonbeek.github.io/industreal.html'),
('IndustReal repository (hand-joint and PSR label details).','https://github.com/TimSchoonbeek/IndustReal'),
('Google MediaPipe Hands documentation: 21 landmarks and normalized coordinates.','https://github.com/google/mediapipe/blob/master/docs/solutions/hands.md'),
('Microsoft HoloLens 2 eye tracking: ~30 Hz gaze ray, calibration, nominal angular accuracy.','https://learn.microsoft.com/en-us/windows/mixed-reality/design/eye-tracking'),
('Microsoft HoloLens 2 Research Mode: depth/IR/IMU streams and frame poses.','https://learn.microsoft.com/en-us/windows/mixed-reality/develop/advanced-concepts/research-mode'),
('Farha & Gall, “MS-TCN: Multi-Stage Temporal Convolutional Network for Action Segmentation,” CVPR 2019.','https://arxiv.org/abs/1903.01945'),
('Li et al., “MS-TCN++,” temporal action segmentation.','https://arxiv.org/abs/2006.09220'),
('Yi et al., “ASFormer: Transformer for Action Segmentation,” BMVC 2021.','https://www.bmva-archive.org.uk/bmvc/2021/assets/papers/0183.pdf'),
('Ego4D Forecasting benchmark: hand/object interaction annotation and sequence tasks.','https://ego4d-data.org/docs/benchmarks/forecasting/'),
('Wen et al., “FoundationPose: Unified 6D Pose Estimation and Tracking of Novel Objects,” CVPR 2024.','https://openaccess.thecvf.com/content/CVPR2024/html/Wen_FoundationPose_Unified_6D_Pose_Estimation_and_Tracking_of_Novel_Objects_CVPR_2024_paper.html'),
('CVAT auto-annotation API: custom model adapter protocol.','https://zhiltsov-max.github.io/cvat/docs/api_sdk/sdk/auto-annotation/'),
('CVAT project repository: SDK, AI-assisted labelling, tracking and QA capabilities.','https://github.com/cvat-ai/cvat')]
for a,u in sources: linktext(a,u)

doc.add_paragraph('Evidence-to-decision traceability', style='Heading 2')
table(['Decision','Evidence used'],[
('Replicate IndustReal-style modalities, not its exact task assumptions','IndustReal paper/project/repository.'),
('Use 21 hand landmarks as the interchange schema','MediaPipe Hands documentation and IndustReal’s joint annotations.'),
('Sensor-first gaze and calibration gate','Microsoft eye tracking and Research Mode documentation.'),
('Small temporal model and segment metrics','MS-TCN/MS-TCN++ and ASFormer action-segmentation literature.'),
('Human review in a model-assisted tool','CVAT auto-annotation interface and tracking workflow.'),
('CAD/reference-based 6D pose is optional','FoundationPose model assumptions and evaluation scope.'),
('Participant/session splits and hand-object context','Ego4D annotation/forecasting design plus standard leakage control.')],[2.5,4.15])

doc.add_paragraph('Final recommendation', style='Heading 1')
callout('Build now', 'MVP = synchronized acquisition + canonical schema + MediaPipe hand proposals + sensor gaze projection + MS-TCN++ step proposals + CVAT review + blind QA. Defer tool orientation until RGB-D calibration and CAD/reference assets are proven. This delivers useful annotation acceleration while keeping scientific claims honest and labels auditable.', GREEN)

doc.save(OUT)
print(OUT.resolve())
